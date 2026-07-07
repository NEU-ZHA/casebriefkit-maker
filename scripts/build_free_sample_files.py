#!/usr/bin/env python3
"""Build the public free CaseBriefKit sample DOCX and PDF files."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = ROOT / "downloads"
DOCX_PATH = DOWNLOADS / "free-case-brief-template.docx"
PDF_PATH = DOWNLOADS / "free-case-brief-template.pdf"


FIELDS = [
    ("Case Name and Citation", "Example: Marbury v. Madison, 5 U.S. 137"),
    ("Court and Date", ""),
    ("Procedural Posture", ""),
    ("Key Facts", ""),
    ("Issue", ""),
    ("Rule", ""),
    ("Analysis", ""),
    ("Holding", ""),
    ("Reasoning", ""),
    ("Class Notes", ""),
]


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_docx(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    for style_name, size, color in [
        ("Heading 1", 18, RGBColor(23, 33, 31)),
        ("Heading 2", 12, RGBColor(31, 111, 91)),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(4)
        style.paragraph_format.space_after = Pt(6)


def build_docx() -> None:
    doc = Document()
    style_docx(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Free Case Brief Template")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(23, 33, 31)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("A one-page law school reading-note form from CaseBriefKit")
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(9)
    subtitle_run.font.color.rgb = RGBColor(93, 104, 100)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run("Educational study aid only. Not legal advice.")
    note_run.font.name = "Arial"
    note_run.font.size = Pt(8)
    note_run.font.color.rgb = RGBColor(93, 104, 100)

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.75)
    table.columns[1].width = Inches(5.0)

    for label, hint in FIELDS:
        row = table.add_row()
        row.cells[0].width = Inches(1.75)
        row.cells[1].width = Inches(5.0)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)
        set_cell_shading(row.cells[0], "EEF5F8")
        label_para = row.cells[0].paragraphs[0]
        label_run = label_para.add_run(label)
        label_run.bold = True
        label_run.font.name = "Arial"
        label_run.font.size = Pt(8.8)
        label_run.font.color.rgb = RGBColor(23, 33, 31)
        value_para = row.cells[1].paragraphs[0]
        value_run = value_para.add_run(hint or " ")
        value_run.font.name = "Arial"
        value_run.font.size = Pt(8.8)
        value_run.font.color.rgb = RGBColor(93, 104, 100)
        if label in {"Key Facts", "Analysis", "Reasoning", "Class Notes"}:
            for _ in range(2):
                row.cells[1].add_paragraph(" ")

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run("casebriefkit-maker")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(93, 104, 100)

    DOWNLOADS.mkdir(exist_ok=True)
    doc.save(DOCX_PATH)


def build_pdf() -> None:
    DOWNLOADS.mkdir(exist_ok=True)
    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.75 * inch,
        leftMargin=0.75 * inch,
        topMargin=0.72 * inch,
        bottomMargin=0.72 * inch,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "Title",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=18,
        leading=22,
        textColor=colors.HexColor("#17211F"),
        alignment=1,
        spaceAfter=4,
    )
    small = ParagraphStyle(
        "Small",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8,
        leading=10,
        textColor=colors.HexColor("#5D6864"),
        alignment=1,
        spaceAfter=8,
    )
    label_style = ParagraphStyle(
        "Label",
        parent=styles["BodyText"],
        fontName="Helvetica-Bold",
        fontSize=8.5,
        leading=10,
        textColor=colors.HexColor("#17211F"),
    )
    value_style = ParagraphStyle(
        "Value",
        parent=styles["BodyText"],
        fontName="Helvetica",
        fontSize=8.5,
        leading=10,
        textColor=colors.HexColor("#5D6864"),
    )

    story = [
        Paragraph("Free Case Brief Template", title_style),
        Paragraph("A one-page law school reading-note form from CaseBriefKit", small),
        Paragraph("Educational study aid only. Not legal advice.", small),
    ]

    rows = []
    for label, hint in FIELDS:
        rows.append([Paragraph(label, label_style), Paragraph(hint or "&nbsp;", value_style)])
    table = Table(rows, colWidths=[1.65 * inch, 5.25 * inch], repeatRows=0)
    table.setStyle(
        TableStyle(
            [
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#D9DFDC")),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#EEF5F8")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
                ("BOTTOMPADDING", (1, 3), (1, 3), 34),
                ("BOTTOMPADDING", (1, 6), (1, 6), 34),
                ("BOTTOMPADDING", (1, 8), (1, 8), 30),
                ("BOTTOMPADDING", (1, 9), (1, 9), 28),
            ]
        )
    )
    story.extend([table, Spacer(1, 8), Paragraph("casebriefkit-maker", small)])
    doc.build(story)


def main() -> None:
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)


if __name__ == "__main__":
    main()
